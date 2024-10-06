# import streamlit as st
# import os
# from typing import List, Dict
# import requests
# from dotenv import load_dotenv
# from sklearn.feature_extraction.text import TfidfVectorizer
# from sklearn.metrics.pairwise import cosine_similarity
# import numpy as np
# from docx import Document

# # Load environment variables
# load_dotenv()

# class VectorDatabase:
#     def __init__(self):
#         self.vectorizer = TfidfVectorizer()
#         self.documents = []
#         self.vectors = None

#     def add_documents(self, documents: List[str]):
#         self.documents = documents
#         self.vectors = self.vectorizer.fit_transform(documents)

#     def search(self, query: str, k: int = 2) -> List[str]:
#         query_vector = self.vectorizer.transform([query])
#         similarities = cosine_similarity(query_vector, self.vectors)[0]
#         top_k_indices = np.argsort(similarities)[-k:][::-1]
#         return [self.documents[i] for i in top_k_indices]

# class OpenAIInterface:
#     def __init__(self, api_key: str):
#         self.api_key = api_key
#         self.base_url = "https://api.openai.com/v1"
#         self.headers = {
#             "Authorization": f"Bearer {self.api_key}",
#             "Content-Type": "application/json"
#         }

#     def generate_chat_response(self, conversation: List[Dict[str, str]]) -> str:
#         payload = {
#             "model": "gpt-3.5-turbo",
#             "messages": conversation
#         }
#         try:
#             response = requests.post(f"{self.base_url}/chat/completions", headers=self.headers, json=payload)
#             response.raise_for_status()  # Will raise HTTPError for bad responses
#             return response.json()['choices'][0]['message']['content']
#         except requests.exceptions.HTTPError as http_err:
#             st.error(f"HTTP error occurred: {http_err}")
#         except Exception as err:
#             st.error(f"An error occurred: {err}")
#         return "Error generating response"

# def chunk_document(text: str, chunk_size: int = 200) -> List[str]:
#     words = text.split()
#     return [' '.join(words[i:i+chunk_size]) for i in range(0, len(words), chunk_size)]

# def read_docx(file):
#     doc = Document(file)
#     return " ".join([paragraph.text for paragraph in doc.paragraphs])

# @st.cache_data
# def process_document(uploaded_file):
#     if uploaded_file.type == "text/plain":
#         return uploaded_file.getvalue().decode("utf-8")
#     elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
#         return read_docx(uploaded_file)
#     return None

# def main():
#     st.set_page_config(page_title="Legal Assistant", page_icon="⚖️", layout="wide")
#     st.title("🤖 Your AI Legal Assistant")

#     # Prompt user to input their OpenAI API Key
#     with st.form(key="api_key_form"):
#         user_api_key = st.text_input("Please enter your OpenAI API Key", type="password")
#         submit_button = st.form_submit_button(label="Submit")

#     if not user_api_key and submit_button:
#         st.error("Please provide a valid OpenAI API Key.")
#         return

#     # Initialize the Vector Database
#     if 'vector_db' not in st.session_state:
#         st.session_state.vector_db = VectorDatabase()

#     # Initialize the OpenAI interface
#     if 'openai_interface' not in st.session_state:
#         st.session_state.openai_interface = OpenAIInterface(user_api_key)

#     if 'document_processed' not in st.session_state:
#         st.session_state.document_processed = False

#     if 'chat_history' not in st.session_state:
#         st.session_state.chat_history = []

#     # File uploader for document
#     uploaded_file = st.file_uploader("Upload a legal document", type=["txt", "docx"])
    
#     if uploaded_file and not st.session_state.document_processed:
#         document = process_document(uploaded_file)
        
#         if document:
#             chunks = chunk_document(document)
#             st.session_state.vector_db.add_documents(chunks)
#             st.session_state.document_processed = True
#             st.success(f"✅ {uploaded_file.name} processed and added to vector database!")
#         else:
#             st.error("Unsupported file type. Please upload a .txt or .docx file.")
#             return

#     # Chat functionality
#     if st.session_state.document_processed:
#         st.subheader("💬 Chat with Your Legal Assistant")
#         for message in st.session_state.chat_history:
#             with st.chat_message(message["role"]):
#                 st.markdown(message["content"])

#         question = st.chat_input("Ask a question about the legal document...")

#         if question:
#             st.session_state.chat_history.append({"role": "user", "content": question})
#             with st.chat_message("user"):
#                 st.markdown(question)

#             with st.spinner("Searching document and generating response..."):
#                 relevant_chunks = st.session_state.vector_db.search(question)
#                 context = "\n".join(relevant_chunks)
#                 prompt = f"Context: {context}\n\nQuestion: {question}\n\nAnswer:"
#                 conversation = [
#                     {"role": "system", "content": "You are a helpful legal assistant. Provide clear and concise answers based on the given context."},
#                     {"role": "user", "content": prompt}
#                 ]
#                 response = st.session_state.openai_interface.generate_chat_response(conversation)

#             if response:
#                 with st.chat_message("assistant"):
#                     st.markdown(response)
#                     st.session_state.chat_history.append({"role": "assistant", "content": response})

# if __name__ == "__main__":
#     main()


import streamlit as st
import os
from typing import List, Dict
import requests
from dotenv import load_dotenv
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np
from docx import Document

load_dotenv()

OPENAI_API_KEY = st.secrets["OPENAI_API_KEY"]
if not OPENAI_API_KEY:
    raise ValueError("OPENAI_API_KEY not found in environment variables")

class VectorDatabase:
    def __init__(self):
        self.vectorizer = TfidfVectorizer()
        self.documents = []
        self.vectors = None

    def add_documents(self, documents: List[str]):
        self.documents = documents
        self.vectors = self.vectorizer.fit_transform(documents)

    def search(self, query: str, k: int = 2) -> List[str]:
        query_vector = self.vectorizer.transform([query])
        similarities = cosine_similarity(query_vector, self.vectors)[0]
        top_k_indices = np.argsort(similarities)[-k:][::-1]
        return [self.documents[i] for i in top_k_indices]

class OpenAIInterface:
    def __init__(self):
        self.api_key = OPENAI_API_KEY
        self.base_url = "https://api.openai.com/v1"
        self.headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }

    def generate_chat_response(self, conversation: List[Dict[str, str]]) -> str:
        payload = {
            "model": "gpt-3.5-turbo",
            "messages": conversation
        }
        response = requests.post(f"{self.base_url}/chat/completions", headers=self.headers, json=payload)
        response.raise_for_status()
        return response.json()['choices'][0]['message']['content']

def chunk_document(text: str, chunk_size: int = 200) -> List[str]:
    words = text.split()
    return [' '.join(words[i:i+chunk_size]) for i in range(0, len(words), chunk_size)]

def read_docx(file):
    doc = Document(file)
    return " ".join([paragraph.text for paragraph in doc.paragraphs])

def main():
    st.set_page_config(page_title="Legal Assistant", page_icon="⚖️", layout="wide")
    st.title("🤖 Your AI Legal Assistant")

    if 'vector_db' not in st.session_state:
        st.session_state.vector_db = VectorDatabase()

    if 'openai_interface' not in st.session_state:
        st.session_state.openai_interface = OpenAIInterface()

    if 'document_processed' not in st.session_state:
        st.session_state.document_processed = False

    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []

    uploaded_file = st.file_uploader("Upload a legal document", type=["txt", "docx"])

    if uploaded_file and not st.session_state.document_processed:
        if uploaded_file.type == "text/plain":
            document = uploaded_file.getvalue().decode("utf-8")
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            document = read_docx(uploaded_file)
        else:
            st.error("Unsupported file type. Please upload a .txt or .docx file.")
            return

        chunks = chunk_document(document)
        st.session_state.vector_db.add_documents(chunks)
        st.session_state.document_processed = True
        st.success(f"✅ {uploaded_file.name} processed and added to vector database!")

    if st.session_state.document_processed:
        st.subheader("💬 Chat with Your Legal Assistant")
        for message in st.session_state.chat_history:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])

        question = st.chat_input("Ask a question about the legal document...")

        if question:
            st.session_state.chat_history.append({"role": "user", "content": question})
            with st.chat_message("user"):
                st.markdown(question)

            with st.spinner("Searching document and generating response..."):
                relevant_chunks = st.session_state.vector_db.search(question)
                context = "\n".join(relevant_chunks)
                prompt = f"Context: {context}\n\nQuestion: {question}\n\nAnswer:"
                conversation = [
                    {"role": "system", "content": "You are a helpful legal assistant. Provide clear and concise answers based on the given context."},
                    {"role": "user", "content": prompt}
                ]
                response = st.session_state.openai_interface.generate_chat_response(conversation)

            with st.chat_message("assistant"):
                st.markdown(response)
                st.session_state.chat_history.append({"role": "assistant", "content": response})

if __name__ == "__main__":
    main()
